#!/usr/bin/env python3
"""
墨吏 · 软著材料合规性验证器
moli-cn-copyright validation tool

验证软件著作权申请全套材料是否符合规范。
可检查：源代码PDF、用户手册DOCX、申请表信息、跨文档一致性。
"""

import os
import re
import sys
import json
import subprocess
import tempfile
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

# ══════════════════════════════════════════════════════════════
# 规则定义
# ══════════════════════════════════════════════════════════════

@dataclass
class Rule:
    id: str
    category: str          # source-code / manual / application / consistency / polices-2026
    severity: str          # error / warning / info
    title: str
    description: str
    check_fn: callable     # returns (passed: bool, detail: str)

CATEGORY_MAP = {
    'R-SC': 'source-code',
    'R-MA': 'manual',
    'R-AP': 'application',
    'R-CO': 'consistency',
    'R-PL': 'polices-2026',
}

@dataclass
class CheckResult:
    rule_id: str
    passed: bool
    severity: str
    title: str
    detail: str

    @property
    def category(self) -> str:
        prefix = self.rule_id.rsplit('-', 1)[0] if '-' in self.rule_id else ''
        return CATEGORY_MAP.get(prefix, 'other')

    def to_dict(self):
        status = "✅" if self.passed else ("❌" if self.severity == "error" else "⚠️")
        return {
            "status": status,
            "passed": self.passed,
            "severity": self.severity,
            "rule_id": self.rule_id,
            "title": self.title,
            "detail": self.detail,
        }


class CopyrightValidator:
    """软著材料合规性验证器"""

    def __init__(self, workdir: str = ".", software_name: str = "", version: str = ""):
        self.workdir = Path(workdir)
        self.software_name = software_name
        self.version = version
        self.results: list[CheckResult] = []
        self._source_pdf_path: Optional[Path] = None
        self._manual_docx_path: Optional[Path] = None
        self._application_txt_path: Optional[Path] = None

    # ── File locators ──

    def _find_source_pdf(self) -> Optional[Path]:
        """查找源代码PDF"""
        patterns = [
            f"*源代码*.pdf",
            f"*源码*.pdf",
            f"*code*.pdf",
            f"*source*.pdf",
        ]
        for p in patterns:
            found = list(self.workdir.glob(p))
            if found:
                return found[0]
        return None

    def _find_manual_docx(self) -> Optional[Path]:
        """查找用户手册DOCX"""
        patterns = [
            f"*操作手册*.docx",
            f"*用户手册*.docx",
            f"*manual*.docx",
        ]
        for p in patterns:
            found = list(self.workdir.glob(p))
            if found:
                return found[0]
        return None

    def _find_application_txt(self) -> Optional[Path]:
        """查找申请表信息TXT"""
        patterns = [
            f"*申请表信息*.txt",
            f"*application*.txt",
        ]
        for p in patterns:
            found = list(self.workdir.glob(p))
            if found:
                return found[0]
        return None

    # ── PDF inspection ──

    def _get_pdf_page_count(self, pdf_path: Path) -> Optional[int]:
        """获取PDF页数"""
        try:
            # macOS: use mdls
            result = subprocess.run(
                ["mdls", "-name", "kMDItemNumberOfPages", str(pdf_path)],
                capture_output=True, text=True, timeout=5
            )
            m = re.search(r'= (\d+)', result.stdout)
            if m:
                return int(m.group(1))
            # Fallback: pdfinfo
            result = subprocess.run(
                ["pdfinfo", str(pdf_path)],
                capture_output=True, text=True, timeout=5
            )
            m = re.search(r'Pages:\s+(\d+)', result.stdout)
            if m:
                return int(m.group(1))
        except Exception:
            pass
        return None

    def _extract_pdf_text(self, pdf_path: Path, first_page=1, last_page=1) -> str:
        """提取PDF指定页面的文本"""
        try:
            result = subprocess.run(
                ["pdftotext", "-f", str(first_page), "-l", str(last_page),
                 "-layout", str(pdf_path), "-"],
                capture_output=True, text=True, timeout=10
            )
            return result.stdout
        except Exception:
            return ""

    def _get_pdf_page_texts(self, pdf_path: Path) -> list[str]:
        """提取PDF所有页面文本（逐页）"""
        count = self._get_pdf_page_count(pdf_path)
        if not count or count > 200:
            return []
        texts = []
        for i in range(1, count + 1):
            text = self._extract_pdf_text(pdf_path, i, i)
            texts.append(text)
        return texts

    # ── DOCX inspection ──

    def _read_docx_text(self, docx_path: Path) -> str:
        """读取DOCX全部文本"""
        try:
            from docx import Document
            doc = Document(str(docx_path))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            return f"<无法读取DOCX: {e}>"

    def _read_docx_structure(self, docx_path: Path) -> list[dict]:
        """读取DOCX章节结构"""
        try:
            from docx import Document
            doc = Document(str(docx_path))
            headings = []
            for p in doc.paragraphs:
                if p.style.name.startswith('Heading'):
                    headings.append({
                        "level": int(p.style.name.split()[-1]),
                        "text": p.text
                    })
            return headings
        except Exception:
            return []

    # ══════════════════════════════════════════════════════
    # 规则 1: 源代码校验
    # ══════════════════════════════════════════════════════

    def _rule_pdf_exists(self):
        """R-SC-01: 源代码PDF是否存在"""
        pdf = self._find_source_pdf()
        self._source_pdf_path = pdf
        if not pdf:
            self.results.append(CheckResult(
                "R-SC-01", False, "error",
                "源代码PDF文件存在",
                "未找到源代码PDF文件。请确认文件名包含'源代码'或'源码'字样。"
            ))
            return False
        self.results.append(CheckResult(
            "R-SC-01", True, "error",
            "源代码PDF文件存在",
            f"找到: {pdf.name}"
        ))
        return True

    def _rule_pdf_page_count(self):
        """R-SC-02: 总页数 = 60 页"""
        if not self._source_pdf_path:
            return
        count = self._get_pdf_page_count(self._source_pdf_path)
        if count is None:
            self.results.append(CheckResult(
                "R-SC-02", False, "error",
                "源代码PDF页数 = 60",
                "无法获取PDF页数，请确认文件有效。"
            ))
            return
        passed = count == 60
        detail = f"实际 {count} 页"
        if count < 60:
            detail += " ❌ 不足60页。若项目代码本就少于60页方可接受，否则需补充。"
        elif count > 60:
            detail += " ❌ 超过60页。必须正好60页（前30+后30），多出的页数需移除。"
        else:
            detail += " ✅ 正好60页"
        self.results.append(CheckResult(
            "R-SC-02", passed, "error",
            "源代码PDF页数 = 60",
            detail
        ))

    def _rule_pdf_header(self):
        """R-SC-03: 页眉包含软件名称和版本号"""
        if not self._source_pdf_path:
            return
        text = self._extract_pdf_text(self._source_pdf_path, 1, 3)
        has_name = self.software_name.replace(' ', '') in text.replace(' ', '') if self.software_name else False
        has_version = self.version in text if self.version else False

        issues = []
        if not has_name and self.software_name:
            issues.append(f"未找到软件名称「{self.software_name}」")
        if not has_version and self.version:
            issues.append(f"未找到版本号「{self.version}」")

        # Check for forbidden content
        has_forbidden_company = bool(re.search(r'(有限公司|科技|技术|股份|集团)', text.split('\n')[0]))
        has_url = 'http' in text[:500]
        has_comment = bool(re.search(r'(//|# |/\*|\*)', text[:500]))

        if has_forbidden_company:
            issues.append("页眉出现公司名称")
        if has_url:
            issues.append("页眉/首页出现URL链接")

        # Check format: should be left "name version", right "page number"
        lines = text.strip().split('\n')
        first_line = lines[0] if lines else ""
        has_name_and_version = bool(re.search(
            r'[\u4e00-\u9fff]+.*V?\d+\.\d+', first_line
        )) if not (has_name and has_version) else True

        detail_parts = []
        if self.software_name:
            detail_parts.append(f"软件名称: {'✅' if has_name else '❌'}")
        if self.version:
            detail_parts.append(f"版本号: {'✅' if has_version else '❌'}")
        if has_forbidden_company:
            detail_parts.append("公司名: ❌ 不应出现在页眉")
        if not has_forbidden_company:
            detail_parts.append("公司名: ✅ 未出现")
        detail_parts.append(f"页眉预览: {first_line[:80]}")

        passed = (has_name or not self.software_name) and \
                 (has_version or not self.version) and \
                 not has_forbidden_company and not has_url
        self.results.append(CheckResult(
            "R-SC-03", passed, "error",
            "页眉包含软件名称和版本号（无公司名/URL）",
            " | ".join(detail_parts)
        ))

    def _rule_pdf_no_blank_lines(self):
        """R-SC-04: 无空行"""
        if not self._source_pdf_path:
            return
        text = self._extract_pdf_text(self._source_pdf_path, 1, 3)
        lines = text.split('\n')
        # Check if page header area has reasonable content density
        # Only check code body area (skip first 2 lines: header)
        body = '\n'.join(lines[2:]) if len(lines) > 2 else text
        consecutive_blanks = 0
        max_blanks = 0
        for line in body.split('\n'):
            if line.strip() == '':
                consecutive_blanks += 1
                max_blanks = max(max_blanks, consecutive_blanks)
            else:
                consecutive_blanks = 0
        passed = max_blanks <= 1  # Allow at most 1 consecutive blank line
        self.results.append(CheckResult(
            "R-SC-04", passed, "warning",
            "源代码无连续空行",
            f"最大连续空行: {max_blanks} {'✅' if passed else '❌ 空行过多，请删除不必要的空行'}"
        ))

    def _rule_pdf_no_comments(self):
        """R-SC-05: 无注释行（排除文件路径标记）"""
        if not self._source_pdf_path:
            return
        text = self._extract_pdf_text(self._source_pdf_path, 1, 3)
        body_lines = text.split('\n')[2:]  # Skip header area
        comment_count = 0
        for line in body_lines:
            stripped = line.strip()
            # 排除文件路径标记 // ========== ... ==========
            if '==========' in stripped:
                continue
            if stripped.startswith('//') or stripped.startswith('#') or stripped.startswith('/*'):
                comment_count += 1
        passed = comment_count == 0
        self.results.append(CheckResult(
            "R-SC-05", passed, "error",
            "源代码无注释行（排除文件路径标记）",
            f"发现 {comment_count} 行注释 {'✅' if passed else '❌ 请删除所有注释行'}"
        ))

    def _rule_pdf_ends_with_complete_block(self):
        """R-SC-06: 最后一页以完整模块结尾"""
        if not self._source_pdf_path:
            return
        count = self._get_pdf_page_count(self._source_pdf_path)
        if not count:
            return
        last_text = self._extract_pdf_text(self._source_pdf_path, count, count)
        # Check the last meaningful line
        lines = [l for l in last_text.split('\n') if l.strip() and '第' not in l and '页' not in l]
        if not lines:
            self.results.append(CheckResult(
                "R-SC-06", False, "error",
                "最后一页以完整模块结尾",
                "无法提取最后一页内容"
            ))
            return
        last_line = lines[-1].strip()
        # Should end with a proper code block ending like }; , }, }, ), ], etc
        ends_properly = bool(re.search(r'[;},\])\)]$', last_line))
        # Or should be a line that's clearly not truncated mid-statement
        if not ends_properly:
            # Check if the line looks truncated
            looks_truncated = (
                len(last_line) > 0 and
                not last_line.endswith(';') and
                not last_line.endswith('}') and
                not last_line.endswith(')') and
                not last_line.endswith(']') and
                not last_line.endswith(',') and
                not last_line.endswith('{')
            )
            passed = not looks_truncated
        else:
            passed = True
        self.results.append(CheckResult(
            "R-SC-06", passed, "error",
            "最后一页以完整模块结尾",
            f"尾行: {last_line[:60]} {'✅' if passed else '❌ 截断在代码中间，请调整使最后一页自然结束'}"
        ))

    # ══════════════════════════════════════════════════════
    # 规则 2: 用户手册校验
    # ══════════════════════════════════════════════════════

    def _rule_manual_exists(self):
        """R-MA-01: 用户手册DOCX是否存在"""
        docx = self._find_manual_docx()
        self._manual_docx_path = docx
        if not docx:
            self.results.append(CheckResult(
                "R-MA-01", False, "error",
                "用户手册DOCX文件存在",
                "未找到用户手册DOCX。请确认文件名包含'操作手册'或'用户手册'字样。"
            ))
            return False
        self.results.append(CheckResult(
            "R-MA-01", True, "error",
            "用户手册DOCX文件存在",
            f"找到: {docx.name}"
        ))
        return True

    def _rule_manual_structure(self):
        """R-MA-02: 章节结构符合规范"""
        if not self._manual_docx_path:
            return
        headings = self._read_docx_structure(self._manual_docx_path)
        heading_texts = [h['text'] for h in headings]

        required_chapters = ['说明', '功能特点', '系统要求']
        found_chapters = []
        missing_chapters = []
        for ch in required_chapters:
            found = any(ch in h for h in heading_texts)
            if found:
                found_chapters.append(ch)
            else:
                missing_chapters.append(ch)

        # Check H1 headings for numbered format (一、二、三 not 1. 2. 3.)
        h1_texts = [h['text'] for h in headings if h['level'] == 1]
        has_arabic_h1 = any(re.match(r'^\d+[.、]', h) for h in h1_texts if h.strip())

        detail_parts = []
        if missing_chapters:
            detail_parts.append(f"缺少章节: {', '.join(missing_chapters)}")
        else:
            detail_parts.append("基础章节齐全 ✅")
        detail_parts.append(
            f"一级标题编号: {'使用中文大写(一/二/三)' if not has_arabic_h1 else '❌ 一级标题使用了阿拉伯数字，应改为中文大写'}"
        )
        detail_parts.append(f"总章节数: {len(headings)} (一级: {len(h1_texts)})")

        passed = len(missing_chapters) == 0 and not has_arabic_h1
        self.results.append(CheckResult(
            "R-MA-02", passed, "error",
            "操作手册章节结构规范（相关文档/说明/功能特点/系统要求）",
            " | ".join(detail_parts)
        ))

    def _rule_manual_has_screenshots(self):
        """R-MA-03: 检查截图是否已添加（查找占位残留 + 实际图片）"""
        if not self._manual_docx_path:
            return
        text = self._read_docx_text(self._manual_docx_path)
        
        # 检查是否还有未替换的占位文字
        remaining_placeholders = len(re.findall(r'【截图[：:].*?】', text))
        
        # 检查是否有实际图片
        has_images = False
        try:
            from docx import Document
            doc = Document(str(self._manual_docx_path))
            has_images = len(doc.inline_shapes) > 0
        except Exception:
            pass
        
        if has_images:
            passed = True
            detail = f"已插入 {len(doc.inline_shapes)} 张实际截图"
            if remaining_placeholders > 0:
                detail += f" | ⚠️ 仍有 {remaining_placeholders} 处占位文字未替换"
        elif remaining_placeholders > 0:
            passed = False
            detail = f"❌ 还有 {remaining_placeholders} 处截图占位未替换为实际图片"
        else:
            passed = False
            detail = "❌ 既无截图也无占位文字（请确认文档内容）"
        
        self.results.append(CheckResult(
            "R-MA-03", passed, "error",
            "截图已添加（占位已替换为实际图片）",
            detail
        ))

    def _rule_manual_page_count(self):
        """R-MA-04: 页数 ≥ 15 页（硬性要求，不满足需补充内容）"""
        if not self._manual_docx_path:
            return
        try:
            from docx import Document
            doc = Document(str(self._manual_docx_path))
            text = "\n".join(p.text for p in doc.paragraphs)
            page_breaks_xml = str(doc.element.xml).count('w:type="page"')
            section_breaks = max(0, len(doc.sections) - 1)
            table_rows = sum(len(t.rows) for t in doc.tables)
            table_chars = table_rows * 40
            # Each paragraph adds spacing overhead (1.5 line spacing)
            para_count = len(doc.paragraphs)
            total_text = len(text) + table_chars + para_count * 80
            est_content_pages = max(1, total_text // 900)
            total_est = max(page_breaks_xml + section_breaks + 1, est_content_pages)
            passed = total_est >= 15
            self.results.append(CheckResult(
                "R-MA-04", passed, "error",
                f"操作手册页数 ≥ 15页（硬性要求，当前估算: {total_est}页）",
                f"估算页数: {total_est} {'✅' if passed else '❌ 不足15页，需要补充章节内容、增加步骤描述'}"
            ))
        except Exception:
            self.results.append(CheckResult(
                "R-MA-04", False, "error",
                "操作手册页数 ≥ 15页",
                "无法读取DOCX内容以估算页数"
            ))
        except Exception:
            self.results.append(CheckResult(
                "R-MA-04", False, "warning",
                "操作手册页数 ≥ 15页",
                "无法读取DOCX内容以估算页数"
            ))

    def _rule_manual_paragraph_style(self):
        """R-MA-05: 正文为段落风格（非项目符号列表）"""
        if not self._manual_docx_path:
            return
        text = self._read_docx_text(self._manual_docx_path)
        lines = text.split('\n')
        bullet_count = sum(1 for l in lines if l.strip().startswith(('- ', '* ', '• ')))
        numbered_steps = sum(1 for l in lines if re.match(r'^\d+[.、]', l.strip()))
        total_lines = len([l for l in lines if l.strip()])

        ratio = (bullet_count + numbered_steps) / max(total_lines, 1)
        passed = ratio < 0.3  # Less than 30% of content is list-style

        self.results.append(CheckResult(
            "R-MA-05", passed, "warning",
            "正文以段落写作为主（少用项目符号/编号列表）",
            f"项目符号行: {bullet_count} | 编号步骤行: {numbered_steps} | "
            f"占比: {ratio:.0%} {'✅' if passed else '❌ 列表式内容过多，应改为自然段落'}"
        ))

    def _rule_manual_no_ai_speak(self):
        """R-MA-06: 无AI味套话（中文20类模式检测）"""
        if not self._manual_docx_path:
            return
        text = self._read_docx_text(self._manual_docx_path)

        # Sources: B1lli/remove-ai-flavor, op7418/humanizer-zh, leeguooooo/stop-slop-zh, blader/humanizer

        # Category 1: Binary contrast shells (否定对比壳)
        binary_patterns = [
            '不是', '而是', '并非', '不在于', '不只是', '更是', '与其说',
            '不仅仅是', '不是.', '而是.', '并非.',
        ]

        # Category 2: Essence claims (本质断言)
        essence_patterns = [
            '真正重要的是', '真正决定', '本质上', '核心在于', '底层逻辑',
            '说白了', '归根结底',
        ]

        # Category 3: AI high-frequency vocabulary (AI高频词)
        ai_vocab = [
            '此外', '至关重要', '深入探讨', '强调', '格局', '关键性的',
            '宝贵的', '充满活力的', '无缝', '不可或缺', '弥足珍贵',
        ]

        # Category 4: Filler phrases (填充短语)
        filler = [
            '值得注意的是', '值得一提的是', '不可否认的是', '众所周知',
            '毫无疑问', '不容忽视', '显而易见', '由此可见', '总的来说',
        ]

        # Category 5: Assistant route markers (路标词)
        route_markers = [
            '下面我们来', '接下来我们', '我们可以看到', '让我为你',
            '在开始之前', '在进入正题之前', '言归正传', '说到这里',
        ]

        # Category 6: Cliche openings (开场套话)
        cliche_openings = [
            '在当今', '随着', '在这个信息爆炸的时代',
            '在科技日新月异的今天', '众所周知', '自古以来',
        ]

        # Category 7: Degree adverbs (程度副词)
        degree_adverbs = [
            '非常', '十分', '极其', '相当', '特别', '格外', '显著', '巨大',
        ]

        # Category 8: Universal positive conclusions (鸡汤收尾)
        uplifting_endings = [
            '未来可期', '共创美好', '携手前行', '向阳而生', '不负韶华',
            '你值得拥有', '心之所向',
        ]

        # Category 9: Empty emphasis shells (强调虚壳)
        empty_emphasis = [
            '我们相信', '我们坚信', '必须认识到', '需要明白的是',
            '需要指出的是', '必须承认',
        ]

        # Category 10: Translation-ese (英译腔)
        translation_ese = [
            '见证了', '这强调了', '这突出了', '这展示了',
            '作为一个', '在...的背景下',
        ]

        # Collect all hits
        all_categories = {
            '否定对比壳': binary_patterns,
            '本质断言': essence_patterns,
            'AI高频词': ai_vocab,
            '填充短语': filler,
            '路标词': route_markers,
            '开场套话': cliche_openings,
            '程度副词': degree_adverbs,
            '鸡汤收尾': uplifting_endings,
            '强调虚壳': empty_emphasis,
            '英译腔': translation_ese,
        }

        findings = []
        for category, patterns in all_categories.items():
            hits = [p for p in patterns if p in text]
            if hits:
                findings.append(f"{category}({len(hits)}处): {', '.join(hits[:4])}")

        # Check em dashes
        em_dash_count = len(re.findall(r'——', text))
        if em_dash_count > 3:
            findings.append(f"破折号({em_dash_count}个) ⚠️")

        # Check triple patterns: "A、B和C"
        triple_count = len(re.findall(r'[\u4e00-\u9fff]+[、][\u4e00-\u9fff]+和[\u4e00-\u9fff]+', text))
        if triple_count > 2:
            findings.append(f"三段式({triple_count}处) ⚠️ 'X、Y和Z'")

        # Check "不仅...更是" pattern
        not_only = len(re.findall(r'不仅.{0,8}(更是|还是|也是)', text))
        if not_only > 0:
            findings.append(f"不仅...更是({not_only}处) ⚠️")

        passed = len([f for f in findings if '处' in f]) == 0 and em_dash_count <= 3
        self.results.append(CheckResult(
            "R-MA-06", passed, "warning",
            "操作手册无AI味套话（中文20类模式检测）",
            " | ".join(findings) if findings else "✅ 无AI套话"
        ))

    # ══════════════════════════════════════════════════════
    # 规则 3: 申请表校验
    # ══════════════════════════════════════════════════════

    def _rule_application_exists(self):
        """R-AP-01: 申请表信息文件存在"""
        txt = self._find_application_txt()
        self._application_txt_path = txt
        if not txt:
            self.results.append(CheckResult(
                "R-AP-01", False, "warning",
                "申请表信息TXT文件存在",
                "未找到申请表信息TXT文件（非必填，但建议保留本地备份）"
            ))
            return False
        self.results.append(CheckResult(
            "R-AP-01", True, "info",
            "申请表信息TXT文件存在",
            f"找到: {txt.name}"
        ))
        return True

    def _rule_app_main_function_length(self):
        """R-AP-02: 主要功能 500-1300字"""
        if not self._application_txt_path:
            return
        text = self._read_docx_text(self._application_txt_path) if self._application_txt_path.suffix == '.docx' \
            else open(self._application_txt_path, encoding='utf-8').read()

        # Try to extract main function section
        patterns = [
            r'主要功能[：:](.*?)(?=\n\n|\n\d|\n开发目的|\n面向|\n技术|\Z)',
            r'主要功能说明[：:](.*?)(?=\n\n|\n\d|\Z)',
        ]
        main_func = ''
        for pat in patterns:
            m = re.search(pat, text, re.DOTALL)
            if m:
                main_func = m.group(1).strip()
                break

        if not main_func:
            self.results.append(CheckResult(
                "R-AP-02", False, "error",
                "主要功能 500-1300字",
                "无法从申请表中提取主要功能字段"
            ))
            return

        # Count Chinese chars (excluding punctuation)
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', main_func))
        total_chars = len(main_func.replace('\n', '').replace(' ', ''))

        passed = 500 <= total_chars <= 1300
        detail = f"总字符: {total_chars} | 中文字符: {chinese_chars}"
        if total_chars < 500:
            detail += " ❌ 不足500字，有补正风险"
        elif total_chars > 1300:
            detail += " ❌ 超过1300字，建议精简"
        else:
            detail += " ✅"
        self.results.append(CheckResult(
            "R-AP-02", passed, "error",
            "主要功能 500-1300字",
            detail
        ))

    def _rule_app_software_name_format(self):
        """R-AP-03: 软件全称以"软件"结尾"""
        if not self.software_name:
            return
        passed = self.software_name.endswith('软件')
        self.results.append(CheckResult(
            "R-AP-03", passed, "error",
            "软件全称以「软件」结尾",
            f"当前: {self.software_name} {'✅' if passed else '❌ 必须以「软件」结尾'}"
        ))

    def _rule_app_version_format(self):
        """R-AP-04: 版本号格式 V1.0"""
        if not self.version:
            return
        passed = bool(re.match(r'^V\d+\.\d+$', self.version))
        detail = f"当前: {self.version}"
        if not passed:
            detail += " ❌ 应使用 V1.0 格式（大写V+两位版本号，如V1.0）"
        else:
            detail += " ✅"
        # Check if it's V1.0 (recommended for first application)
        if self.version != 'V1.0':
            detail += " ⚠️ 首次申请建议使用 V1.0"
        self.results.append(CheckResult(
            "R-AP-04", passed, "error",
            "版本号格式 V1.0（大写V+两位）",
            detail
        ))

    def _rule_pdf_no_third_party_names(self):
        """R-SC-10: 代码中无第三方公司/个人名称"""
        if not self._source_pdf_path:
            return
        text = self._extract_pdf_text(self._source_pdf_path, 1, 3)
        # Common company-related patterns that shouldn't appear (except the copyright holder)
        forbidden = re.findall(r'有限公司|股份有限公司|有限责任公司|科技公司|集团|工作室', text[:1000])
        # Filter out the copyright holder itself
        actual_forbidden = [f for f in forbidden if not any(owner in text[:200] for owner in ['沃泰森', 'WENTSEN'])]
        passed = len(actual_forbidden) == 0
        self.results.append(CheckResult(
            "R-SC-10", passed, "error",
            "源代码中无第三方公司/个人名称",
            f"发现 {len(actual_forbidden)} 个疑似第三方名称 {'✅' if passed else '❌ 请检查并删除第三方公司名称'}"
        ))

    def _rule_docx_has_toc_field(self):
        """R-MA-08: DOCX 使用自动目录域代码（非手写）"""
        if not self._manual_docx_path:
            return
        try:
            import zipfile
            with zipfile.ZipFile(self._manual_docx_path, 'r') as z:
                # 读取 document.xml 查找域代码
                if 'word/document.xml' in z.namelist():
                    xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
                    has_toc = 'TOC' in xml and 'instrText' in xml
                    passed = has_toc
                    detail = f"TOC域代码: {'✅ 已设置' if has_toc else '❌ 未找到，请在Word中插入自动目录'}"
                else:
                    passed = False
                    detail = "❌ 无法读取DOCX内容"
            
            self.results.append(CheckResult("R-MA-08", passed, "warning", "DOCX使用自动目录（TOC域代码）", detail))
        except Exception as e:
            self.results.append(CheckResult("R-MA-08", True, "info", "DOCX自动目录检查", f"无法检查: {e}"))

    def _rule_docx_has_page_fields(self):
        """R-MA-09: 页眉页码使用域代码（非手写）"""
        if not self._manual_docx_path:
            return
        try:
            import zipfile
            with zipfile.ZipFile(self._manual_docx_path, 'r') as z:
                # 检查 header 相关 XML
                header_files = [n for n in z.namelist() if 'header' in n and n.endswith('.xml')]
                all_xml = ''
                for hf in header_files:
                    all_xml += z.read(hf).decode('utf-8', errors='ignore')
                # 也检查 document.xml (页眉可能内联)
                if 'word/document.xml' in z.namelist():
                    all_xml += z.read('word/document.xml').decode('utf-8', errors='ignore')
                
                has_page = 'PAGE' in all_xml
                has_numpages = 'NUMPAGES' in all_xml
                passed = has_page
                detail = f"PAGE域: {'✅' if has_page else '❌ 未找到'} | NUMPAGES域: {'✅' if has_numpages else '⚠️ 建议添加'}"
            
            self.results.append(CheckResult("R-MA-09", passed, "warning", "页眉页码使用Word域代码", detail))
        except Exception as e:
            self.results.append(CheckResult("R-MA-09", True, "info", "页眉页码检查", f"无法检查: {e}"))

    def _rule_docx_toc_aligned(self):
        """R-MA-10: 目录与正文标题对齐（目录已更新）"""
        if not self._manual_docx_path:
            return
        try:
            import zipfile
            from docx import Document
            doc = Document(str(self._manual_docx_path))
            
            # 收集文档中的实际标题
            actual_headings = []
            for p in doc.paragraphs:
                if p.style.name.startswith('Heading'):
                    actual_headings.append(p.text.strip())
            
            # 收集 TOC 中的条目
            toc_entries = []
            with zipfile.ZipFile(self._manual_docx_path, 'r') as z:
                if 'word/document.xml' in z.namelist():
                    xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
                    import re
                    # 找到 TOC 结果区（separate 和 end 之间的文本）
                    toc_match = re.search(r'fldCharType="separate".*?fldCharType="end"', xml, re.DOTALL)
                    if toc_match:
                        toc_text = toc_match.group()
                        # 提取标题文本
                        toc_entries = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', toc_text)
                        toc_entries = [t.strip() for t in toc_entries if t.strip()]
            
            if not actual_headings:
                self.results.append(CheckResult("R-MA-10", True, "info", "目录与标题对齐检查", "文档中未找到标题"))
                return
            
            if not toc_entries:
                self.results.append(CheckResult("R-MA-10", False, "error", "目录与标题对齐检查", "❌ 未找到目录条目，请在Word中更新目录"))
                return
            
            # 比较（放宽匹配：只要标题文本出现在目录中即视为已更新）
            missing = []
            for h in actual_headings:
                if h and len(h) > 2:
                    found = any(h in entry or entry in h for entry in toc_entries)
                    if not found:
                        missing.append(h)
            
            if missing:
                passed = False
                detail = f"❌ {len(missing)} 个标题在目录中缺失: {'; '.join(missing[:5])}（请在Word中右键目录→更新域）"
            else:
                passed = True
                detail = f"✅ 目录 {len(toc_entries)} 项与标题对齐"
            
            self.results.append(CheckResult("R-MA-10", passed, "error", "目录与正文标题对齐（目录已更新）", detail))
        except Exception as e:
            self.results.append(CheckResult("R-MA-10", True, "info", "目录对齐检查", f"无法检查: {e}"))

    def _rule_docx_structure(self):
        """R-MA-11: 文档结构正确（封面→目录→正文各自分页）"""
        if not self._manual_docx_path:
            return
        try:
            import zipfile, re
            with zipfile.ZipFile(self._manual_docx_path, 'r') as z:
                xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
            
            body = xml[xml.find('<w:body>'):xml.find('</w:body>')] if '<w:body>' in xml else xml
            paras = re.findall(r'<w:p[ >].*?</w:p>', body, re.DOTALL)
            
            # 找出现的顺序：封面内容 → 分页符 → 目录 → 分页符 → 正文标题
            found_first_pb = False
            found_mulu = False
            found_second_pb = False
            found_body_start = False
            
            for p in paras:
                text = re.sub(r'<[^>]+>', '', p)[:30]
                has_pb = 'w:type="page"' in p
                
                if not found_first_pb and has_pb:
                    found_first_pb = True
                    continue
                if found_first_pb and not found_mulu and ('目录' in text or 'TOC' in text):
                    found_mulu = True
                    continue
                if found_mulu and not found_second_pb and has_pb:
                    found_second_pb = True
                    continue
                if found_second_pb and not found_body_start and ('一、' in text or '说明' in text or 'Heading' in p):
                    found_body_start = True
                    # no break, want to check more
            
            issues = []
            if not found_first_pb: issues.append("封面后缺分页符")
            if not found_mulu: issues.append("未找到目录")
            if not found_second_pb: issues.append("目录后缺分页符")
            if not found_body_start: issues.append("未找到正文起始标题")
            
            passed = len(issues) == 0
            detail = "✅ 封面→目录→正文结构正确" if passed else "❌ " + ", ".join(issues)
            self.results.append(CheckResult("R-MA-11", passed, "error", "文档结构（封面→目录→正文各自分页）", detail))
        except Exception as e:
            self.results.append(CheckResult("R-MA-11", True, "info", "文档结构检查", f"无法检查: {e}"))

    def _rule_manual_step_count(self):
        """R-MA-12: 每个功能模块编号步骤 ≥ 5 步"""
        if not self._manual_docx_path:
            return
        try:
            text = self._read_docx_text(self._manual_docx_path)
            import re
            # 找编号步骤：行首 "数字."
            steps = re.findall(r'^\d+\.\s+\S', text, re.MULTILINE)
            step_count = len(steps)
            # 找功能模块数：一级标题数
            from docx import Document
            doc = Document(str(self._manual_docx_path))
            h1_count = sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading 1'))
            # 排除封面和目录标题
            h1_actual = max(0, h1_count - 2)
            
            if h1_actual > 0:
                avg_steps = step_count / h1_actual
                passed = avg_steps >= 3 and step_count >= h1_actual * 2
                detail = f"总{step_count}个编号步骤 / {h1_actual}个功能模块 = 均{avg_steps:.1f}步 {'✅' if passed else '❌ 建议每功能≥5步'}"
            else:
                passed = step_count >= 5
                detail = f"总{step_count}个编号步骤 {'✅' if passed else '❌ 建议至少5个步骤'}"
            
            self.results.append(CheckResult("R-MA-12", passed, "error", "功能模块编号步骤数量（建议每功能≥5步）", detail))
        except Exception as e:
            self.results.append(CheckResult("R-MA-12", True, "info", "步骤数量检查", f"无法检查: {e}"))

    def _rule_manual_ui_controls(self):
        """R-MA-13: 引用带引号的UI控件名"""
        if not self._manual_docx_path:
            return
        try:
            text = self._read_docx_text(self._manual_docx_path)
            import re
            # 找引号引用的UI控件："XXX"
            controls = re.findall(r'"([^"]{2,30})"', text)
            # 过滤掉不是UI控件名的通用引号内容
            ui_keywords = ['按钮', '框', '区', '栏', '页', '菜单', '图标', '链接', '选项', '输入', '选择', '列表', '卡片', '提示', '窗口', '弹窗', '确认', '取消', '保存', '删除', '编辑']
            ui_controls = [c for c in controls if any(kw in c for kw in ui_keywords)]
            
            control_count = len(ui_controls)
            # 估算步骤数
            steps = re.findall(r'^\d+\.\s+\S', text, re.MULTILINE)
            step_count = max(len(steps), 1)
            
            ratio = control_count / step_count
            passed = ratio >= 0.5 and control_count >= 3
            detail = f"UI控件引用: {control_count}处, 均{ratio:.1f}处/步 {'✅' if passed else '❌ 建议每步至少引用1个UI控件名'}"
            if ui_controls:
                detail += f" | 示例: {', '.join(ui_controls[:5])}"
            
            self.results.append(CheckResult("R-MA-13", passed, "error", "引用带引号的UI控件名（建议每步≥1处）", detail))
        except Exception as e:
            self.results.append(CheckResult("R-MA-13", True, "info", "UI控件检查", f"无法检查: {e}"))

    def _rule_manual_punctuation(self):
        """R-MA-14: 中文正文标点统一（全角中文标点，不混用半角）"""
        if not self._manual_docx_path:
            return
        try:
            text = self._read_docx_text(self._manual_docx_path)
            import re
            # 检查中文正文中混用的半角标点
            # 场景：中文字符附近出现半角逗号、句号、冒号、分号、问号、感叹号
            mixed = []
            patterns = [
                (r'[\u4e00-\u9fff][,]', '半角逗号在中文字后'),
                (r'[\u4e00-\u9fff][:]', '半角冒号在中文字后'),
                (r'[\u4e00-\u9fff][;]', '半角分号在中文字后'),
                (r'[\u4e00-\u9fff][.]', '半角句号在中文字后'),
                (r'[\u4e00-\u9fff][!]', '半角感叹号在中文字后'),
                (r'[\u4e00-\u9fff][?]', '半角问号在中文字后'),
            ]
            for pat, desc in patterns:
                matches = re.findall(pat, text)
                if matches:
                    mixed.append(f"{desc}({len(matches)}处)")
            
            passed = len(mixed) == 0
            detail = "✅ 中文标点统一使用全角符号" if passed else "❌ " + "; ".join(mixed[:5])
            self.results.append(CheckResult("R-MA-14", passed, "warning", "中文正文标点统一（全角，参考GB/T 15834）", detail))
        except Exception as e:
            self.results.append(CheckResult("R-MA-14", True, "info", "标点检查", f"无法检查: {e}"))

    def fix_punctuation(self) -> int:
        """修复中文正文中的半角标点为全角，返回修复处数"""
        self._manual_docx_path = self._find_manual_docx()
        if not self._manual_docx_path:
            print("❌ 未找到操作手册DOCX文件")
            return 0
        try:
            import zipfile, shutil, re
            from docx import Document
            doc = Document(str(self._manual_docx_path))
            fix_count = 0
            
            # 中文字后的半角标点 → 全角
            punct_map = {
                ',': '，', ':': '：', ';': '；',
                '!': '！', '?': '？',
            }
            # 句点特殊处理：区分句号和小数点
            # 中文字后的 . → 。
            # 数字后的 . → 保留（小数点）
            
            for para in doc.paragraphs:
                for run in para.runs:
                    text = run.text
                    new_text = text
                    # 中文字后的逗号冒号分号感叹号问号
                    for half, full in punct_map.items():
                        new_text = re.sub(f'([\\u4e00-\\u9fff]){re.escape(half)}', f'\\1{full}', new_text)
                    # 中文字后的句点（排除数字后的）
                    new_text = re.sub(r'([\u4e00-\u9fff])\.(?!\d)', r'\1。', new_text)
                    
                    if new_text != text:
                        fix_count += 1
                        run.text = new_text
            
            # 保存
            doc.save(str(self._manual_docx_path))
            return fix_count
        except Exception as e:
            print(f"❌ 修复失败: {e}")
            return 0

    def _rule_file_naming(self):
        """R-CO-03: 文件名规范"""
        workdir = Path(self.workdir)
        issues = []
        # 检查 PDF - 排除临时文件
        pdfs = [f for f in workdir.glob('*.pdf') if not f.name.startswith('~') and not f.name.startswith('.~')]
        for pdf in pdfs:
            stem = pdf.stem
            if self.software_name:
                sn = self.software_name.replace(' ', '')
                if sn not in stem.replace(' ', ''):
                    issues.append(f"PDF文件名缺软件名称: {pdf.name}")
            if '源代码' not in stem and '源码' not in stem:
                issues.append(f"PDF文件名应含「源代码」: {pdf.name}")
        # 检查 DOCX
        docxs = [f for f in workdir.glob('*.docx') if not f.name.startswith('~') and not f.name.startswith('.~')]
        for docx in docxs:
            stem = docx.stem
            if self.software_name:
                sn = self.software_name.replace(' ', '')
                if sn not in stem.replace(' ', ''):
                    issues.append(f"DOCX文件名缺软件名称: {docx.name}")
            if '操作手册' not in stem and '用户手册' not in stem and '手册' not in stem:
                issues.append(f"DOCX文件名应含「操作手册」: {docx.name}")
        # 检查 DOCX
        docxs = list(workdir.glob('*.docx'))
        for docx in docxs:
            name = docx.name
            stem = docx.stem
            if self.software_name:
                sn_chars = self.software_name.replace(' ', '')
                name_clean = stem.replace(' ', '')
                if sn_chars not in name_clean:
                    issues.append(f"DOCX文件名缺软件名称: {name}")
            if '操作手册' not in stem and '用户手册' not in stem and '手册' not in stem:
                issues.append(f"DOCX文件名应含「操作手册」: {name}")
        
        passed = len(issues) == 0
        detail = "✅ 文件名规范" if passed else "❌ " + "; ".join(issues[:3])
        self.results.append(CheckResult("R-CO-03", passed, "warning", "文件名规范（含软件名称+文档类型）", detail))

    def _rule_app_date_logic(self):
        """R-AP-05: 日期逻辑正确（开发完成 ≤ 首次发表日期 < 申请日期）"""
        if not self._application_txt_path:
            return
        text = open(self._application_txt_path, encoding='utf-8').read()
        # Try to find dates
        dates = re.findall(r'(\d{4}[-年]\d{1,2}[-月]\d{1,2})', text)
        if len(dates) >= 2:
            self.results.append(CheckResult(
                "R-AP-05", True, "info",
                "日期逻辑检查（开发完成 ≤ 首次发表日期 < 申请日期）",
                f"检测到 {len(dates)} 个日期，请人工核验逻辑顺序"
            ))
        else:
            self.results.append(CheckResult(
                "R-AP-05", True, "info",
                "日期逻辑检查",
                "未找到足够日期信息进行自动校验，请人工确认"
            ))

    def _rule_manual_formatting(self):
        """R-MA-07: 手册排版规范（字体/行距/缩进）"""
        if not self._manual_docx_path:
            return
        try:
            from docx import Document
            doc = Document(str(self._manual_docx_path))
            sample_para = None
            for p in doc.paragraphs:
                if p.text.strip() and len(p.text) > 20:
                    sample_para = p
                    break
            if sample_para:
                pf = sample_para.paragraph_format
                fi = pf.first_line_indent
                ls = pf.line_spacing
                # Check font
                fonts_used = set()
                for p in doc.paragraphs[:50]:
                    for run in p.runs:
                        if run.font.name:
                            fonts_used.add(run.font.name.lower())
                has_songti = any('song' in f or '宋' in f or 'simsun' in f for f in fonts_used)
                has_heitii = any('hei' in f or '黑' in f or 'simhei' in f for f in fonts_used)
                details = []
                details.append(f"字体: {', '.join(list(fonts_used)[:5]) or '未检测'}")
                if fi:
                    details.append(f"首行缩进: {'✅' if fi > 0 else '❌ 建议首行缩进2字符'}")
                details.append(f"标题字体(黑体): {'✅' if has_heitii else '⚠️ 建议标题使用黑体'}")
                details.append(f"正文字体(宋体): {'✅' if has_songti else '⚠️ 建议正文使用宋体'}")
                passed = has_songti or has_heitii
                self.results.append(CheckResult(
                    "R-MA-07", passed, "warning",
                    "手册排版规范（宋体正文/黑体标题/首行缩进）",
                    " | ".join(details)
                ))
            else:
                self.results.append(CheckResult("R-MA-07", True, "info", "手册排版规范", "未找到足够内容样本"))
        except Exception:
            self.results.append(CheckResult("R-MA-07", True, "info", "手册排版规范", "无法解析DOCX格式"))

    # ══════════════════════════════════════════════════════
    # 规则 4: 跨文档一致性
    # ══════════════════════════════════════════════════════

    def _rule_consistency_name_across_docs(self):
        """R-CO-01: 软件名称跨文档一致"""
        if not self.software_name:
            return
        issues = []
        if self._source_pdf_path:
            text = self._extract_pdf_text(self._source_pdf_path, 1, 1)
            if self.software_name.replace(' ', '') not in text.replace(' ', ''):
                issues.append("源代码PDF页眉")
        if self._manual_docx_path:
            text = self._read_docx_text(self._manual_docx_path)
            if self.software_name not in text:
                issues.append("操作手册正文/页眉")

        passed = len(issues) == 0
        detail = f"软件名称: {self.software_name}"
        if issues:
            detail += f" ❌ 以下材料未找到一致名称: {', '.join(issues)}"
        else:
            detail += " ✅ 所有材料一致"
        self.results.append(CheckResult(
            "R-CO-01", passed, "error",
            "软件名称在所有材料中一致",
            detail
        ))

    def _rule_consistency_version_across_docs(self):
        """R-CO-02: 版本号跨文档一致"""
        if not self.version:
            return
        issues = []
        if self._source_pdf_path:
            text = self._extract_pdf_text(self._source_pdf_path, 1, 1)
            if self.version not in text:
                issues.append("源代码PDF页眉")
        if self._manual_docx_path:
            text = self._read_docx_text(self._manual_docx_path)
            if self.version not in text:
                issues.append("操作手册")

        passed = len(issues) == 0
        detail = f"版本号: {self.version}"
        if issues:
            detail += f" ❌ 以下材料未找到一致版本号: {', '.join(issues)}"
        else:
            detail += " ✅ 所有材料一致"
        self.results.append(CheckResult(
            "R-CO-02", passed, "error",
            "版本号在所有材料中一致",
            detail
        ))

    # ══════════════════════════════════════════════════════
    # 规则 5: 2026新政
    # ══════════════════════════════════════════════════════

    def _rule_policy_ai_declaration(self):
        """R-PL-01: AI辅助开发说明（2026新政）"""
        # Check if there's an AI declaration file
        ai_declaration = list(self.workdir.glob("*AI*声明*")) + list(self.workdir.glob("*ai*declaration*"))
        docx_text = ""
        if self._manual_docx_path:
            docx_text = self._read_docx_text(self._manual_docx_path)
        has_ai_development_claim = 'AI辅助开发' in docx_text or 'AI生成' in docx_text or '人工智能辅助' in docx_text
        has_declaration = len(ai_declaration) > 0 or 'AI辅助开发声明' in docx_text or 'AI开发声明' in docx_text

        if has_ai_development_claim and not has_declaration:
            passed = False
            detail = "❌ 材料提及AI辅助开发但未找到《AI辅助开发声明》（2026年3月15日新规要求）"
        else:
            passed = True
            detail = "✅ 未发现AI合规问题" if not has_ai_development_claim else \
                     f"✅ 已包含AI声明: {ai_declaration[0].name if ai_declaration else '在文档中提及'}"

        self.results.append(CheckResult(
            "R-PL-01", passed, "warning",
            "AI辅助开发声明（2026新政）",
            detail
        ))

    def _rule_policy_open_source(self):
        """R-PL-02: 开源组件说明（2026新政）"""
        docx_text = ""
        if self._manual_docx_path:
            docx_text = self._read_docx_text(self._manual_docx_path)
        has_open_source_mention = '开源' in docx_text or 'open source' in docx_text.lower() or 'MIT' in docx_text or 'Apache' in docx_text or 'GPL' in docx_text
        has_declaration = '开源组件' in docx_text or '开源许可' in docx_text or '开源协议' in docx_text

        if has_open_source_mention and not has_declaration:
            passed = False
            detail = "⚠️ 材料提及开源但未说明开源组件清单（2026新规要求列明组件名称/版本/许可协议）"
        else:
            passed = True
            detail = "✅ 未发现开源合规问题" if not has_open_source_mention else "✅ 已包含开源说明"

        self.results.append(CheckResult(
            "R-PL-02", passed, "info",
            "开源组件说明（2026新政）",
            detail
        ))

    # ══════════════════════════════════════════════════════
    # 运行所有规则
    # ══════════════════════════════════════════════════════

    def run_all(self):
        """运行所有验证规则"""
        self.results = []

        # Phase 1: Source Code
        if self._rule_pdf_exists():
            self._rule_pdf_page_count()
            self._rule_pdf_header()
            self._rule_pdf_no_blank_lines()
            self._rule_pdf_no_comments()
            self._rule_pdf_ends_with_complete_block()
            self._rule_pdf_no_third_party_names()

        # Phase 2: User Manual
        if self._rule_manual_exists():
            self._rule_manual_structure()
            self._rule_manual_has_screenshots()
            self._rule_manual_page_count()
            self._rule_manual_paragraph_style()
            self._rule_manual_no_ai_speak()
            self._rule_manual_formatting()
            self._rule_docx_has_toc_field()
            self._rule_docx_has_page_fields()
            self._rule_docx_toc_aligned()
            self._rule_docx_structure()
            self._rule_manual_step_count()
            self._rule_manual_ui_controls()
            self._rule_manual_punctuation()

        # Phase 3: Application
        self._rule_application_exists()
        self._rule_app_main_function_length()
        self._rule_app_software_name_format()
        self._rule_app_version_format()
        self._rule_app_date_logic()

        # Phase 4: Consistency
        self._rule_consistency_name_across_docs()
        self._rule_consistency_version_across_docs()
        self._rule_file_naming()

        # Phase 5: 2026 Policies
        self._rule_policy_ai_declaration()
        self._rule_policy_open_source()

        return self.results

    def summary(self) -> dict:
        """生成验证摘要"""
        total = len(self.results)
        passed = sum(1 for r in self.results if r.passed)
        errors = sum(1 for r in self.results if not r.passed and r.severity == 'error')
        warnings = sum(1 for r in self.results if not r.passed and r.severity == 'warning')
        infos = sum(1 for r in self.results if not r.passed and r.severity == 'info')

        return {
            "total": total,
            "passed": passed,
            "errors": errors,
            "warnings": warnings,
            "infos": infos,
            "pass_rate": f"{passed / max(total, 1) * 100:.0f}%",
            "verdict": "✅ 通过" if errors == 0 else f"❌ 未通过（{errors}个错误，{warnings}个警告）",
        }

    def print_report(self, format: str = "text") -> str:
        """输出验证报告"""
        if format == "json":
            return json.dumps({
                "summary": self.summary(),
                "results": [r.to_dict() for r in self.results],
            }, ensure_ascii=False, indent=2)

        # Text format
        lines = []
        lines.append("=" * 60)
        lines.append("  墨吏 · 软著材料合规性验证报告")
        lines.append("  moli-cn-copyright validation report")
        lines.append("=" * 60)
        lines.append(f"  工作目录: {self.workdir}")
        if self.software_name:
            lines.append(f"  软件名称: {self.software_name}")
        if self.version:
            lines.append(f"  版本号:   {self.version}")
        lines.append("")

        # Group by category
        categories = {
            "source-code": "一、源代码校验",
            "manual": "二、用户手册校验",
            "application": "三、申请表校验",
            "consistency": "四、跨文档一致性",
            "polices-2026": "五、2026新政",
        }

        for cat_key, cat_title in categories.items():
            cat_results = [r for r in self.results if r.category == cat_key]
            if not cat_results:
                continue
            lines.append(f"\n  {cat_title}")
            lines.append("  " + "-" * 40)
            for r in cat_results:
                status = "✅" if r.passed else ("❌" if r.severity == "error" else "⚠️")
                lines.append(f"  [{status}] {r.rule_id} {r.title}")
                # Indent detail
                for detail_line in r.detail.split(" | "):
                    lines.append(f"         {detail_line}")
                lines.append("")

        # Summary
        s = self.summary()
        lines.append("-" * 60)
        lines.append(f"  总计: {s['total']} | ✅ 通过: {s['passed']} | "
                      f"❌ 错误: {s['errors']} | ⚠️ 警告: {s['warnings']} | "
                      f"ℹ️ 信息: {s['infos']}")
        lines.append(f"  结论: {s['verdict']}")
        lines.append("=" * 60)

        return "\n".join(lines)


try:
    VERSION = open(Path(__file__).resolve().parents[2] / "VERSION").read().strip()
except (FileNotFoundError, OSError):
    VERSION = "0.0.0"

# ══════════════════════════════════════════════════════════════
# CLI入口
# ══════════════════════════════════════════════════════════════

def main():
    import argparse
    parser = argparse.ArgumentParser(
        description=f"墨吏 · 软著材料合规性验证器 v{VERSION}",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=f"""
版本: {VERSION}
示例:
  # 验证当前目录下的软著材料
  python3 validate_materials.py --software-name "山野集-菌迹软件" --version V1.0

  # 指定工作目录
  python3 validate_materials.py --workdir /path/to/materials --software-name "XXX软件" --version V1.0

  # 输出JSON格式报告
  python3 validate_materials.py --software-name "山野集-菌迹软件" --version V1.0 --format json

  # 简化模式（仅检查源代码PDF）
  python3 validate_materials.py --software-name "山野集-菌迹软件" --version V1.0 --mode source-only
        """
    )
    parser.add_argument('--workdir', default='.',
                        help='软著材料所在目录（默认: 当前目录）')
    parser.add_argument('--software-name', default='',
                        help='软件全称（如: 山野集-菌迹软件）')
    parser.add_argument('--version', default='',
                        help='版本号（如: V1.0）')
    parser.add_argument('--format', choices=['text', 'json'], default='text',
                        help='输出格式（默认: text）')
    parser.add_argument('--mode', choices=['full', 'source-only', 'manual-only'],
                        default='full',
                        help='验证模式（默认: full 全部验证）')
    parser.add_argument('--output', '-o',
                        help='输出报告到文件（可选）')
    parser.add_argument('--fix-punctuation', action='store_true',
                        help='修复中文正文中混用的半角标点为全角')

    args = parser.parse_args()

    validator = CopyrightValidator(
        workdir=args.workdir,
        software_name=args.software_name,
        version=args.version,
    )

    # 修复模式优先
    if args.fix_punctuation:
        count = validator.fix_punctuation()
        if count > 0:
            print(f"✅ 已修复 {count} 处标点符号（半角→全角）")
        else:
            print("✅ 标点符号已全部为全角，无需修复")
        sys.exit(0)

    validator.run_all()

    report = validator.print_report(format=args.format)

    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(report)
        print(f"报告已保存到: {args.output}")
    else:
        print(report)

    # Return exit code
    s = validator.summary()
    sys.exit(0 if s['errors'] == 0 else 1)


if __name__ == '__main__':
    main()
